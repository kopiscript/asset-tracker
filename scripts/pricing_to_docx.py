from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Cm(2.0)
section.bottom_margin = Cm(2.0)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)

# ── Colour palette ────────────────────────────────────────────────────────────
TEAL   = RGBColor(0x00, 0xC2, 0xCC)   # brand accent
DARK   = RGBColor(0x1A, 0x1A, 0x2E)   # near-black
MUTED  = RGBColor(0x6B, 0x72, 0x80)   # grey
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT  = RGBColor(0xF3, 0xF4, 0xF6)   # table header bg
BORDER = RGBColor(0xE5, 0xE7, 0xEB)

# ── Helpers ───────────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def set_cell_border(cell, border_color="E5E7EB"):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), border_color)
        tcBorders.append(el)
    tcPr.append(tcBorders)

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"

    # header row
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        set_cell_bg(cell, "00C2CC")
        set_cell_border(cell, "00C2CC")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = WHITE
        cell._tc.get_or_add_tcPr()

    # data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        bg  = "FFFFFF" if r_idx % 2 == 0 else "F9FAFB"
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            set_cell_bg(cell, bg)
            set_cell_border(cell)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            if str(val).startswith("**") and str(val).endswith("**"):
                run.text = val[2:-2]
                run.bold = True

    # column widths
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    return table

def h1(doc, text):
    p   = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size  = Pt(22)
    run.font.bold  = True
    run.font.color.rgb = DARK
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(4)
    # accent underline rule
    pPr   = p._p.get_or_add_pPr()
    pBdr  = OxmlElement("w:pBdr")
    btm   = OxmlElement("w:bottom")
    btm.set(qn("w:val"),   "single")
    btm.set(qn("w:sz"),    "6")
    btm.set(qn("w:space"), "4")
    btm.set(qn("w:color"), "00C2CC")
    pBdr.append(btm)
    pPr.append(pBdr)

def h2(doc, text):
    p   = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size  = Pt(13)
    run.font.bold  = True
    run.font.color.rgb = TEAL
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)

def h3(doc, text):
    p   = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size  = Pt(11)
    run.font.bold  = True
    run.font.color.rgb = DARK
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)

def body(doc, text, italic=False, color=None):
    p   = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size   = Pt(10)
    run.font.italic = italic
    run.font.color.rgb = color if color else DARK
    p.paragraph_format.space_after = Pt(4)
    return p

def note(doc, text):
    p   = doc.add_paragraph()
    run = p.add_run("ℹ  " + text)
    run.font.size   = Pt(9)
    run.font.italic = True
    run.font.color.rgb = MUTED
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.space_after  = Pt(6)

def spacer(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)

# ═══════════════════════════════════════════════════════════════════════════════
# DOCUMENT
# ═══════════════════════════════════════════════════════════════════════════════

h1(doc, "Mirae — Internal Cost, Pricing & Margin Model")

meta = doc.add_paragraph()
for label, val in [("Last updated: ", "2026-05-25"), ("  |  Status: ", "Working draft"), ("  |  Audience: ", "Internal — AZP Group only")]:
    r = meta.add_run(label)
    r.font.size = Pt(9)
    r.font.color.rgb = MUTED
    r2 = meta.add_run(val)
    r2.font.size = Pt(9)
    r2.font.bold = True
    r2.font.color.rgb = DARK
meta.paragraph_format.space_after = Pt(12)

# ─── 1. Infrastructure OpEx ───────────────────────────────────────────────────
h2(doc, "1. Infrastructure OpEx  (Vercel + Neon + Resend only)")

h3(doc, "Data growth rate")
add_table(doc,
    ["Metric", "Value"],
    [
        ["Ping frequency", "1 per minute per vehicle  =  1,440 rows / day / vehicle"],
        ["Row size estimate", "~500 bytes  (30 telemetry fields + index overhead)"],
        ["Growth per vehicle", "~21 MB / month   |   ~252 MB / year"],
    ],
    col_widths=[5, 11]
)
spacer(doc)

h3(doc, "Storage at scale  (6-month rolling retention)")
add_table(doc,
    ["Vehicles", "Storage used", "Neon plan", "Neon cost / mo"],
    [
        ["≤ 80",        "≤ 9.8 GB",        "Launch  (10 GB)",       "$19"],
        ["81 – 400",    "9.9 – 49.2 GB",   "Scale  (50 GB)",        "$69"],
        ["401 – 600",   "49.3 – 73.8 GB",  "Scale + overage",       "$69 + ~$3–4"],
        ["601 – 1,000", "73.9 – 123 GB",   "Scale + overage",       "$69 + ~$4–11"],
        ["1,000+",      "> 123 GB",         "Scale + overage",       "~$80–100"],
    ],
    col_widths=[3, 4.5, 4.5, 4]
)
note(doc, "Neon overage rate: $0.15 / GB-month — very cheap above the 50 GB base.")

h3(doc, "Monthly infrastructure cost by scale")
add_table(doc,
    ["Vehicles", "Neon", "Vercel Pro", "Resend", "Total USD / mo", "Total MYR / mo"],
    [
        ["≤ 80",         "$19", "$20", "$0 (free)",  "$39",      "~RM 183"],
        ["81 – 400",     "$69", "$20", "$0",          "$89",      "~RM 418"],
        ["401 – 1,000",  "~$73","$20", "$20",         "~$113",    "~RM 531"],
        ["1,000+",       "~$80–100","$20","$20",      "~$120–140","~RM 564–658"],
    ],
    col_widths=[3, 2.5, 2.5, 2.5, 3.5, 3.5]
)
note(doc, "USD/MYR rate used: 4.70. Update periodically.")

# ─── 2. Personnel ─────────────────────────────────────────────────────────────
h2(doc, "2. Personnel OpEx  (Mirae team — 3 people)")

add_table(doc,
    ["Role", "Gross salary", "EPF employer 13%", "SOCSO + EIS ~2%", "All-in cost / mo"],
    [
        ["Developer 1",     "RM ___", "RM ___", "RM ___", "RM ___"],
        ["Developer 2",     "RM ___", "RM ___", "RM ___", "RM ___"],
        ["Developer / PM",  "RM ___", "RM ___", "RM ___", "RM ___"],
        ["Total",           "",        "",        "",        "RM ___"],
    ],
    col_widths=[4, 3.5, 3.5, 3.5, 3.5]
)
note(doc, "Fill in actual salaries. Example: Gross RM 5,000 → EPF RM 650 → SOCSO+EIS RM 100 → All-in RM 5,750.")
note(doc, "Reference total used in this model until filled: RM 19,550 / month.")

# ─── 3. Total cost & cost per vehicle ─────────────────────────────────────────
h2(doc, "3. Total Monthly Cost & Cost per Vehicle")
note(doc, "Formula: Total cost = Personnel + Infrastructure")

add_table(doc,
    ["Vehicles", "Infra (MYR)", "Personnel (MYR)", "Total cost / mo", "Cost per vehicle"],
    [
        ["50",     "RM 183",  "RM 19,550", "RM 19,733", "RM 394.66"],
        ["100",    "RM 418",  "RM 19,550", "RM 19,968", "RM 199.68"],
        ["200",    "RM 418",  "RM 19,550", "RM 19,968", "RM 99.84"],
        ["300",    "RM 418",  "RM 19,550", "RM 19,968", "RM 66.56"],
        ["500",    "RM 531",  "RM 19,550", "RM 20,081", "RM 40.16"],
        ["1,000",  "RM 658",  "RM 19,550", "RM 20,208", "RM 20.21"],
    ],
    col_widths=[3, 3.5, 3.5, 3.5, 3.5]
)
note(doc, "Personnel is the dominant fixed cost. Infra barely moves even at 10× vehicle scale.")
note(doc, "The business is loss-making below ~300 vehicles at market rates. Acquiring clients fast matters more than optimising infra.")

# ─── 4. Break-even & margin ────────────────────────────────────────────────────
h2(doc, "4. Break-even & Margin Analysis")
note(doc, "Formula: Revenue = Total cost ÷ (1 − target margin)")

h3(doc, "At 300 vehicles  (target launch milestone)")
add_table(doc,
    ["Target margin", "Revenue needed / mo", "Revenue per vehicle / mo"],
    [
        ["0%  (break-even)", "RM 19,968",  "RM 66.56"],
        ["20%",              "RM 24,960",  "RM 83.20"],
        ["30%",              "RM 28,526",  "RM 95.09"],
        ["40%",              "RM 33,280",  "RM 110.93"],
        ["50%",              "RM 39,936",  "RM 133.12"],
    ],
    col_widths=[4.5, 5.5, 6]
)
spacer(doc)

h3(doc, "At 500 vehicles")
add_table(doc,
    ["Target margin", "Revenue needed / mo", "Revenue per vehicle / mo"],
    [
        ["0%  (break-even)", "RM 20,081",  "RM 40.16"],
        ["20%",              "RM 25,101",  "RM 50.20"],
        ["30%",              "RM 28,687",  "RM 57.37"],
        ["40%",              "RM 33,468",  "RM 66.94"],
        ["50%",              "RM 40,162",  "RM 80.32"],
    ],
    col_widths=[4.5, 5.5, 6]
)
note(doc, "Market rate for fleet tracking SaaS in Malaysia: RM 30–80 / vehicle / month.")
note(doc, "Pricing at RM 60–70/vehicle puts you at ~40% margin at 500 vehicles — competitive and healthy.")

# ─── 5. Suggested pricing ──────────────────────────────────────────────────────
h2(doc, "5. Suggested Pricing Structure  (to be finalised)")

h3(doc, "Option A — Per-vehicle flat rate")
add_table(doc,
    ["Tier", "Vehicles / org", "Data retention", "Price / vehicle / mo", "Notes"],
    [
        ["Starter",    "Up to 15",    "3 months",  "RM ___", "Small fleets, low commitment"],
        ["Standard",   "Up to 50",    "6 months",  "RM ___", "Typical dealership"],
        ["Enterprise", "Unlimited",   "12 months", "RM ___", "Multi-branch, custom SLA"],
    ],
    col_widths=[3, 3.5, 3.5, 4, 4]
)
spacer(doc)

h3(doc, "Option B — Platform fee + per-vehicle")
body(doc, "Base fee:    RM 200 / org / month   (covers platform access)")
body(doc, "Per vehicle: RM 45–60 / vehicle / month")
spacer(doc)
body(doc, "Example:  org with 25 vehicles  →  RM 200 + (25 × RM 50)  =  RM 1,450 / month")
body(doc, "At 15 such orgs  →  RM 21,750 / month  →  profitable above 300 vehicles")

# ─── 6. Data retention lever ──────────────────────────────────────────────────
h2(doc, "6. Data Retention as a Pricing Lever")
body(doc, "Retention period is the single biggest cost control knob AND a tier differentiator.")

add_table(doc,
    ["Retention", "Storage per vehicle", "Neon plan for 300 vehicles"],
    [
        ["3 months",  "~63 MB",   "Launch  ($19)"],
        ["6 months",  "~126 MB",  "Scale   ($69)"],
        ["12 months", "~252 MB",  "Scale   ($69) + small overage"],
    ],
    col_widths=[4, 5, 7]
)
note(doc, "Recommendation: 6 months as default. Offer 12 months as a paid tier add-on.")
note(doc, "Retention is enforced by the nightly cron at POST /api/cron/retention (configurable via RETENTION_MONTHS env var).")

# ─── 7. Update log ────────────────────────────────────────────────────────────
h2(doc, "7. Assumptions & Update Log")
add_table(doc,
    ["Date", "Update"],
    [
        ["2026-05-25", "Initial model created. Personnel figures are estimates — replace with actuals."],
        ["",           ""],
        ["",           ""],
    ],
    col_widths=[4, 13]
)

# ── Save ──────────────────────────────────────────────────────────────────────
out = "/home/user/asset-tracker/docs/Mirae_Pricing_Model.docx"
doc.save(out)
print(f"Saved: {out}")
